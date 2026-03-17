from __future__ import annotations

from io import BytesIO


def build_cam_docx(cam_text: str) -> bytes | None:
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    doc.add_heading("Intelli-Credit: Credit Appraisal Memo", level=1)
    for line in cam_text.split("\n"):
        if not line.strip():
            doc.add_paragraph("")
        elif line[:2].isdigit() and ". " in line:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_cam_pdf(cam_text: str) -> bytes | None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception:
        return None

    bio = BytesIO()
    c = canvas.Canvas(bio, pagesize=A4)
    width, height = A4

    y = height - 40
    c.setFont("Helvetica-Bold", 14)
    c.drawString(40, y, "Intelli-Credit: Credit Appraisal Memo")
    y -= 24

    c.setFont("Helvetica", 9)
    for raw_line in cam_text.split("\n"):
        line = raw_line.strip()
        if not line:
            y -= 10
        else:
            while line:
                chunk = line[:110]
                line = line[110:]
                c.drawString(40, y, chunk)
                y -= 11
                if y < 40:
                    c.showPage()
                    c.setFont("Helvetica", 9)
                    y = height - 40
        if y < 40:
            c.showPage()
            c.setFont("Helvetica", 9)
            y = height - 40

    c.save()
    bio.seek(0)
    return bio.getvalue()
