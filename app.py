from __future__ import annotations

from io import BytesIO

import pandas as pd
import streamlit as st

from modules import (
    build_cam_text,
    build_cam_docx,
    build_cam_pdf,
    build_indian_context_cards,
    build_local_rag_index,
    build_research_agent_findings,
    combine_scorecard_ml,
    compute_five_cs_score,
    derive_rag_adjustments,
    engineer_features,
    extract_document_signals,
    extract_pdf_text_with_fallback,
    extract_qualitative_adjustments,
    fetch_live_research,
    generate_credit_recommendation,
    generate_research_insights,
    get_demo_scenarios,
    load_annual_report_text,
    load_bank_statement,
    load_gst_summary,
    load_manual_notes,
    load_structured_or_pdf,
    reconcile_gst_bank,
    retrieve_evidence_by_theme,
    summarize_source_text,
    summarize_structured_source,
    train_hybrid_model,
    predict_ml_view,
    MODEL_VERSION,
)

st.set_page_config(page_title="Intelli-Credit", page_icon="IC", layout="wide")


def _status_chip(label: str, value: str) -> str:
    color = "#0f766e"
    if str(value).lower() in {"review", "moderate", "watch", "possible", "placeholder"}:
        color = "#b45309"
    if str(value).lower() in {"reject", "high", "high mismatch", "alert", "failed", "elevated"}:
        color = "#b91c1c"
    return f"<div class='chip'><span>{label}</span><strong style='color:{color}'>{value}</strong></div>"


def _build_docx_bytes(cam_text: str):
    try:
        from docx import Document
    except Exception:
        return None

    doc = Document()
    for line in cam_text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        elif line[:2].isdigit() and ". " in line:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def _read_text_upload(uploaded_file) -> str:
    if uploaded_file is None:
        return ""
    try:
        return uploaded_file.getvalue().decode("utf-8", errors="ignore").strip()
    except Exception:
        return ""


def _source_findings_line(summary: dict) -> str:
    pairs = [f"{k}={v}" for k, v in summary.items() if k != "source"]
    return "; ".join(pairs[:6]) if pairs else "No finding"


@st.cache_resource(show_spinner=False)
def _get_hybrid_model(_version: str):
    return train_hybrid_model()


st.markdown(
    """
    <style>
    .main {background: linear-gradient(180deg, #f7f9fc 0%, #edf3fb 100%);}
    .title-banner {
        background: linear-gradient(120deg, #0c2d57 0%, #144f8f 60%, #1f6fba 100%);
        border-radius: 14px; padding: 18px 24px; color: white; margin-bottom: 14px;
        box-shadow: 0 8px 24px rgba(12, 45, 87, 0.25);
    }
    .summary-banner {
        background: #ffffff; border: 1px solid #dbe6f5; border-radius: 12px; padding: 12px 16px;
        box-shadow: 0 4px 14px rgba(20, 46, 88, 0.08); margin-bottom: 12px;
    }
    .chip {
        display:flex; justify-content:space-between; align-items:center;
        background:#f8fbff; border:1px solid #d8e6f8; border-radius:10px; padding:9px 10px; margin-bottom:8px;
    }
    .flow-box {
        border:1px solid #d1deee; border-radius:10px; background:#ffffff; padding:14px; text-align:center;
        box-shadow: 0 3px 10px rgba(20,46,88,0.06);
        min-height:68px;
    }
    .flow-arrow {text-align:center; font-size:24px; color:#1f6fba; margin-top:12px;}
    </style>
    """,
    unsafe_allow_html=True,
)

st.markdown(
    """
    <div class='title-banner'>
        <h2 style='margin:0;'>Intelli-Credit: AI-Powered Corporate Credit Appraisal</h2>
        <p style='margin:6px 0 0 0;'>Hackathon-ready prototype for Indian corporate lending workflows</p>
    </div>
    """,
    unsafe_allow_html=True,
)

scenarios = get_demo_scenarios()

with st.sidebar:
    st.header("Input Control")
    demo_mode = st.toggle("Demo Mode", value=True)

    selected_scenario = None
    if demo_mode:
        selected_scenario = st.selectbox("Demo Scenario", list(scenarios.keys()), index=0)

    company_name = st.text_input("Company Name", value="Apex Industrial Components Pvt Ltd", disabled=demo_mode)
    sector = st.selectbox(
        "Sector",
        ["Manufacturing", "Infrastructure", "Retail", "Pharma", "IT Services"],
        index=0,
        disabled=demo_mode,
    )

    st.markdown("### Structured")
    bank_file = st.file_uploader("Bank Statement (CSV)", type=["csv"], key="bank", disabled=demo_mode)
    gst_file = st.file_uploader("GST Summary (CSV/XLSX)", type=["csv", "xlsx", "xls"], key="gst", disabled=demo_mode)
    itr_file = st.file_uploader("ITR (CSV/XLSX/PDF)", type=["csv", "xlsx", "xls", "pdf"], key="itr", disabled=demo_mode)
    shareholding_file = st.file_uploader(
        "Shareholding Pattern (CSV/XLSX/PDF)", type=["csv", "xlsx", "xls", "pdf"], key="share", disabled=demo_mode
    )

    st.markdown("### Unstructured")
    annual_report_file = st.file_uploader("Annual Report (PDF)", type=["pdf"], key="annual", disabled=demo_mode)
    legal_notice_file = st.file_uploader("Legal Notice (PDF)", type=["pdf"], key="legal", disabled=demo_mode)
    sanction_letter_file = st.file_uploader("Sanction Letter (PDF)", type=["pdf"], key="sanction", disabled=demo_mode)
    rating_report_file = st.file_uploader("Rating Report (PDF)", type=["pdf"], key="rating", disabled=demo_mode)
    board_minutes_file = st.file_uploader("Board Meeting Minutes (PDF)", type=["pdf"], key="board", disabled=demo_mode)

    st.markdown("### Primary Insights")
    primary_notes_input = st.text_area(
        "Primary Due Diligence Notes",
        placeholder="Add site visit and management quality insights...",
        height=110,
        disabled=demo_mode,
    )

    st.markdown("### External Intelligence")
    external_notes_input = st.text_area(
        "External Intelligence Notes",
        placeholder="Add market/regulatory/promoter intelligence...",
        height=100,
        disabled=demo_mode,
    )
    external_notes_file = st.file_uploader("External Notes File (.txt)", type=["txt"], key="ext", disabled=demo_mode)

    st.markdown("### Research Agent Inputs")
    promoter_obs_input = st.text_area("Promoter / Management Observation", height=70)
    sector_obs_input = st.text_area("Sector Observation", height=70)
    regulatory_obs_input = st.text_area("Regulatory Observation", height=70)
    mca_notes_input = st.text_area("MCA Notes", height=70)
    litigation_notes_input = st.text_area("Litigation Notes", height=70)

    st.markdown("### Optional Live Research")
    live_research_mode = st.toggle("Live Research Mode", value=False)

    run_clicked = st.button("Run Intelli-Credit Analysis", type="primary")

if "ran_once" not in st.session_state:
    st.session_state.ran_once = False
if run_clicked:
    st.session_state.ran_once = True

if not st.session_state.ran_once:
    st.info("Click 'Run Intelli-Credit Analysis' for latest inputs. Demo Mode preloads curated borrower cases.")

source_rows = []
source_findings = {}
unstructured_processing = []

if demo_mode:
    scenario = scenarios[selected_scenario]
    company_name = scenario["company_name"]
    sector = scenario["sector"]

    bank_df = scenario["bank_df"].copy()
    gst_df = scenario["gst_df"].copy()
    primary_notes = scenario["manual_notes"]
    external_notes = scenario.get("external_notes", "")

    annual_result = {
        "text": scenario["report_text"],
        "uploaded": False,
        "parsed": True,
        "method": "demo_text",
        "message": "Demo scenario narrative",
        "chars": len(scenario["report_text"]),
    }

    additional_unstructured = {
        "Legal Notice": {"text": "", "uploaded": False, "parsed": False, "method": "not_provided", "message": "Not provided", "chars": 0},
        "Sanction Letter": {"text": "", "uploaded": False, "parsed": False, "method": "not_provided", "message": "Not provided", "chars": 0},
        "Rating Report": {"text": "", "uploaded": False, "parsed": False, "method": "not_provided", "message": "Not provided", "chars": 0},
        "Board Minutes": {"text": "", "uploaded": False, "parsed": False, "method": "not_provided", "message": "Not provided", "chars": 0},
    }

    itr_result = {"kind": "missing", "dataframe": pd.DataFrame(), "text": "", "parsed": False, "method": "not_provided"}
    share_result = {"kind": "missing", "dataframe": pd.DataFrame(), "text": "", "parsed": False, "method": "not_provided"}

    completeness_note = "High confidence: full scenario dataset available (demo mode)."
    data_source_note = f"Demo Mode: {selected_scenario}"
else:
    bank_df = load_bank_statement(bank_file)
    gst_df = load_gst_summary(gst_file)

    primary_notes = load_manual_notes(primary_notes_input)
    external_file_text = _read_text_upload(external_notes_file)
    external_notes = (external_notes_input or "").strip()
    if external_file_text:
        external_notes = f"{external_notes}\n{external_file_text}".strip()

    annual_fallback = load_annual_report_text(None)
    annual_result = extract_pdf_text_with_fallback(annual_report_file, fallback_text=annual_fallback)

    additional_unstructured = {
        "Legal Notice": extract_pdf_text_with_fallback(legal_notice_file, fallback_text=""),
        "Sanction Letter": extract_pdf_text_with_fallback(sanction_letter_file, fallback_text=""),
        "Rating Report": extract_pdf_text_with_fallback(rating_report_file, fallback_text=""),
        "Board Minutes": extract_pdf_text_with_fallback(board_minutes_file, fallback_text=""),
    }

    itr_result = load_structured_or_pdf(itr_file)
    share_result = load_structured_or_pdf(shareholding_file)

    tracked = [
        annual_report_file,
        bank_file,
        gst_file,
        itr_file,
        shareholding_file,
        legal_notice_file,
        sanction_letter_file,
        rating_report_file,
        board_minutes_file,
    ]
    available_count = sum(1 for item in tracked if item is not None)
    available_count += 1 if primary_notes_input.strip() else 0
    available_count += 1 if external_notes else 0
    completeness_score = int((available_count / 11) * 100)
    completeness_note = f"Data completeness: {completeness_score}% based on provided multi-source inputs."
    data_source_note = "Uploaded files + fallback defaults for missing sources"

combined_unstructured_text_parts = []
if annual_result.get("text"):
    combined_unstructured_text_parts.append(f"[Annual Report] {annual_result['text']}")
for src, result in additional_unstructured.items():
    if result.get("text"):
        combined_unstructured_text_parts.append(f"[{src}] {result['text']}")

if itr_result.get("kind") == "pdf_text" and itr_result.get("text"):
    combined_unstructured_text_parts.append(f"[ITR PDF] {itr_result['text']}")
if share_result.get("kind") == "pdf_text" and share_result.get("text"):
    combined_unstructured_text_parts.append(f"[Shareholding PDF] {share_result['text']}")

combined_unstructured_text = "\n".join(combined_unstructured_text_parts)
combined_notes = f"{primary_notes}\n{external_notes}".strip()

initial_doc_signals = extract_document_signals(combined_unstructured_text or annual_result.get("text", ""))
recon = reconcile_gst_bank(gst_df, bank_df)
insights = generate_research_insights(sector=sector, document_signals=initial_doc_signals, manual_notes=combined_notes)

rag_documents = [
    {"source": "Annual Report", "source_type": "PDF", "source_category": "Unstructured", "text": annual_result.get("text", "")},
]
for src, result in additional_unstructured.items():
    rag_documents.append(
        {
            "source": src,
            "source_type": "PDF",
            "source_category": "Unstructured",
            "text": result.get("text", ""),
        }
    )
if itr_result.get("kind") == "pdf_text":
    rag_documents.append({"source": "ITR PDF", "source_type": "PDF", "source_category": "Structured/Semi", "text": itr_result.get("text", "")})
if share_result.get("kind") == "pdf_text":
    rag_documents.append(
        {
            "source": "Shareholding PDF",
            "source_type": "PDF",
            "source_category": "Structured/Semi",
            "text": share_result.get("text", ""),
        }
    )

rag_index = build_local_rag_index(rag_documents)
rag_theme_results = retrieve_evidence_by_theme(rag_index, top_k=3)
rag_adjustments = derive_rag_adjustments(rag_theme_results)

qual_adjustments = extract_qualitative_adjustments(primary_notes=primary_notes, external_notes=external_notes)
research_agent_inputs = {
    "promoter_obs": promoter_obs_input,
    "sector_obs": sector_obs_input,
    "regulatory_obs": regulatory_obs_input,
    "mca_notes": mca_notes_input,
    "litigation_notes": litigation_notes_input,
}
research_agent_findings = build_research_agent_findings(
    manual_inputs=research_agent_inputs,
    base_insights=insights,
    rag_theme_results=rag_theme_results,
    primary_notes=primary_notes,
    external_notes=external_notes,
)
live_research = fetch_live_research(company_name, sector) if live_research_mode else {
    "mode": "fallback",
    "status": "disabled",
    "headlines": [],
    "summary": "Live research mode disabled; using manual/mock research.",
}
if live_research.get("headlines"):
    top_titles = "; ".join(h.get("title", "")[:90] for h in live_research["headlines"][:3])
    research_agent_findings["cards"]["Live Research Headlines"] = top_titles
    research_agent_findings["summary"] = (
        research_agent_findings.get("summary", "") + f"\\nLive research: {live_research.get('summary', '')}"
    ).strip()

research_agent_adjustments = [
    {
        "source": row.get("source", "Research Agent"),
        "source_type": row.get("origin", "external_intelligence"),
        "signal": row.get("cue", "Research cue"),
        "factor": row.get("five_c", "Conditions"),
        "points": row.get("score_effect", 0),
        "reason": row.get("reason", "Research panel-derived signal"),
    }
    for row in research_agent_findings.get("signal_rows", [])
]

all_adjustments = qual_adjustments + rag_adjustments + research_agent_adjustments

score_bundle = compute_five_cs_score(
    doc_signals=initial_doc_signals,
    recon=recon,
    insights=insights,
    manual_notes=combined_notes,
    qualitative_adjustments=all_adjustments,
)
recommendation = generate_credit_recommendation(score_bundle=score_bundle, doc_signals=initial_doc_signals, recon=recon)

source_rows.append(
    {
        "Source": "Annual Report",
        "Type": "Unstructured",
        "Uploaded": "Yes" if annual_result["uploaded"] else "No",
        "Parsed": "Yes" if annual_result["parsed"] else "No",
        "Used in Analysis": "Yes" if bool(annual_result.get("text")) else "No",
        "Parse Method": annual_result.get("method", "NA"),
    }
)
for src, result in additional_unstructured.items():
    source_rows.append(
        {
            "Source": src,
            "Type": "Unstructured",
            "Uploaded": "Yes" if result.get("uploaded") else "No",
            "Parsed": "Yes" if result.get("parsed") else "No",
            "Used in Analysis": "Yes" if bool(result.get("text")) else "No",
            "Parse Method": result.get("method", "NA"),
        }
    )

source_rows.extend(
    [
        {
            "Source": "Bank Statement",
            "Type": "Structured",
            "Uploaded": "Yes" if bank_file is not None and not demo_mode else "No",
            "Parsed": "Yes",
            "Used in Analysis": "Yes",
            "Parse Method": "csv",
        },
        {
            "Source": "GST Summary",
            "Type": "Structured",
            "Uploaded": "Yes" if gst_file is not None and not demo_mode else "No",
            "Parsed": "Yes",
            "Used in Analysis": "Yes",
            "Parse Method": "csv/excel",
        },
        {
            "Source": "ITR",
            "Type": "Structured/Semi-Structured",
            "Uploaded": "Yes" if itr_file is not None and not demo_mode else "No",
            "Parsed": "Yes" if itr_result.get("parsed") else "No",
            "Used in Analysis": "Yes" if itr_result.get("parsed") else "No",
            "Parse Method": itr_result.get("method", "NA"),
        },
        {
            "Source": "Shareholding Pattern",
            "Type": "Structured/Semi-Structured",
            "Uploaded": "Yes" if shareholding_file is not None and not demo_mode else "No",
            "Parsed": "Yes" if share_result.get("parsed") else "No",
            "Used in Analysis": "Yes" if share_result.get("parsed") else "No",
            "Parse Method": share_result.get("method", "NA"),
        },
        {
            "Source": "Primary Due Diligence Notes",
            "Type": "Primary Insights",
            "Uploaded": "Yes" if primary_notes else "No",
            "Parsed": "Yes" if primary_notes else "No",
            "Used in Analysis": "Yes" if primary_notes else "No",
            "Parse Method": "text",
        },
        {
            "Source": "External Intelligence Notes",
            "Type": "External Intelligence",
            "Uploaded": "Yes" if external_notes else "No",
            "Parsed": "Yes" if external_notes else "No",
            "Used in Analysis": "Yes" if external_notes else "No",
            "Parse Method": "text/txt",
        },
    ]
)

annual_summary = summarize_source_text("Annual Report", annual_result.get("text", ""))
source_findings["Annual Report Findings"] = _source_findings_line(annual_summary)
for src, result in additional_unstructured.items():
    summary = summarize_source_text(src, result.get("text", ""))
    source_findings[f"{src} Findings"] = _source_findings_line(summary)

if itr_result.get("kind") == "dataframe":
    itr_summary = summarize_structured_source("ITR", itr_result.get("dataframe"))
elif itr_result.get("kind") == "pdf_text":
    itr_summary = summarize_source_text("ITR", itr_result.get("text", ""))
else:
    itr_summary = {"source": "ITR", "insight": "No ITR source provided"}
source_findings["ITR Findings"] = _source_findings_line(itr_summary)

if share_result.get("kind") == "dataframe":
    share_summary = summarize_structured_source("Shareholding Pattern", share_result.get("dataframe"))
elif share_result.get("kind") == "pdf_text":
    share_summary = summarize_source_text("Shareholding Pattern", share_result.get("text", ""))
else:
    share_summary = {"source": "Shareholding Pattern", "insight": "No shareholding source provided"}
source_findings["Shareholding Pattern Findings"] = _source_findings_line(share_summary)

source_findings["Primary Due Diligence Findings"] = _source_findings_line(
    summarize_source_text("Primary Notes", primary_notes)
)
source_findings["External Intelligence Findings"] = _source_findings_line(
    summarize_source_text("External Notes", external_notes)
)

unstructured_processing.append(
    {
        "Source": "Annual Report",
        "Uploaded": annual_result.get("uploaded"),
        "Parsed": annual_result.get("parsed"),
        "Method": annual_result.get("method"),
        "Chars": annual_result.get("chars", 0),
        "Message": annual_result.get("message", ""),
    }
)
for src, result in additional_unstructured.items():
    unstructured_processing.append(
        {
            "Source": src,
            "Uploaded": result.get("uploaded"),
            "Parsed": result.get("parsed"),
            "Method": result.get("method"),
            "Chars": result.get("chars", 0),
            "Message": result.get("message", ""),
        }
    )

indian_context_cards = build_indian_context_cards(
    recon=recon,
    insights=insights,
    doc_signals=initial_doc_signals,
    shareholding_summary=share_summary if isinstance(share_summary, dict) else {},
    combined_notes=combined_notes,
)

model_bundle = _get_hybrid_model(MODEL_VERSION)
ml_features = engineer_features(
    score_bundle=score_bundle,
    recon=recon,
    doc_signals=initial_doc_signals,
    insights=insights,
    qualitative_adjustments=qual_adjustments,
    rag_adjustments=rag_adjustments,
    source_rows=source_rows,
    indian_context_cards=indian_context_cards,
)
ml_view = predict_ml_view(model_bundle, ml_features)
hybrid_view = combine_scorecard_ml(
    scorecard_reco=recommendation,
    ml_view=ml_view,
    doc_signals=initial_doc_signals,
    qualitative_adjustments=qual_adjustments,
    rag_adjustments=rag_adjustments,
)

cam_payload = {
    "company_name": company_name,
    "sector": sector,
    "doc_signals": initial_doc_signals,
    "recon": recon,
    "insights": insights,
    "score": score_bundle,
    "recommendation": recommendation,
    "manual_notes": primary_notes,
    "primary_notes": primary_notes,
    "external_notes": external_notes,
    "completeness_note": completeness_note,
    "data_source_note": data_source_note,
    "source_coverage": source_rows,
    "source_findings": source_findings,
    "qualitative_adjustments": all_adjustments,
    "rag_theme_results": rag_theme_results,
    "research_agent": research_agent_findings,
    "indian_context_cards": indian_context_cards,
    "ml_view": ml_view,
    "hybrid_view": hybrid_view,
    "live_research": live_research,
}
cam_text = build_cam_text(cam_payload)
docx_bytes = build_cam_docx(cam_text)
pdf_bytes = build_cam_pdf(cam_text)

st.markdown(
    f"""
    <div class='summary-banner'>
        <h4 style='margin:0 0 8px 0;'>Executive Summary</h4>
        <div style='display:grid;grid-template-columns:repeat(6,1fr);gap:8px;'>
            <div><small>Company</small><br><strong>{company_name}</strong></div>
            <div><small>Sector</small><br><strong>{sector}</strong></div>
            <div><small>Risk Category</small><br><strong>{score_bundle['risk_category']}</strong></div>
            <div><small>Hybrid Decision</small><br><strong>{hybrid_view['final_decision']}</strong></div>
            <div><small>Hybrid Limit</small><br><strong>INR {hybrid_view['hybrid_limit']:,.0f}</strong></div>
            <div><small>Hybrid Pricing</small><br><strong>{hybrid_view['hybrid_pricing_band']}</strong></div>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

k1, k2, k3, k4 = st.columns(4)
k1.metric("Overall Score", f"{score_bundle['total_score']}/100")
k2.metric("Reconciliation Status", recon["overall_status"])
k3.metric("Hybrid Decision", hybrid_view["final_decision"])
k4.metric("Decision Confidence", f"{hybrid_view['decision_confidence']:.2f}")

st.markdown("---")

tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10, tab11, tab12, tab13, tab14 = st.tabs(
    [
        "Screen 1: Inputs & Demo",
        "Screen 2: Document Insights",
        "Screen 3: GST-Bank Reconciliation",
        "Screen 4: Research & Qualitative",
        "Screen 5: Scoring & Explainability",
        "Screen 6: CAM & Recommendation",
        "Unstructured Intelligence",
        "Research Agent",
        "Indian Context Signals",
        "Hybrid Recommendation",
        "ML Diagnostics",
        "Export CAM",
        "Live Research",
        "How It Works",
    ]
)

with tab1:
    st.subheader("Input / Upload / Demo Selection")
    s1, s2 = st.columns(2)
    with s1:
        st.markdown("**Current Input Mode**")
        st.write("Demo Mode" if demo_mode else "Manual Upload Mode")
        if demo_mode:
            st.success(f"Scenario loaded: {selected_scenario}")
        st.caption(data_source_note)
    with s2:
        st.markdown("**Data Completeness & Confidence**")
        st.info(completeness_note)

    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Bank Statement Snapshot**")
        st.dataframe(bank_df.head(12), use_container_width=True)
    with c2:
        st.markdown("**GST Summary Snapshot**")
        st.dataframe(gst_df.head(12), use_container_width=True)

    st.markdown("**Source Coverage Summary**")
    st.dataframe(pd.DataFrame(source_rows), use_container_width=True)

with tab2:
    st.subheader("Extracted Financial and Multi-Document Insights")
    d1, d2, d3, d4 = st.columns(4)
    d1.metric("Revenue", f"INR {initial_doc_signals['revenue']:,.0f}")
    d2.metric("Debt", f"INR {initial_doc_signals['debt']:,.0f}")
    d3.metric("Liabilities", f"INR {initial_doc_signals['liabilities']:,.0f}")
    d4.metric("Operating Cash Flow", f"INR {initial_doc_signals['operating_cash_flow']:,.0f}")

    st.write(f"Cash Flow Indicator: **{initial_doc_signals['cash_flow_indicator']}**")

    x1, x2 = st.columns(2)
    with x1:
        st.markdown("**Litigation / Legal Notice Indicator**")
        st.markdown(_status_chip("Litigation Signal", str(initial_doc_signals["litigation_mentions"])), unsafe_allow_html=True)
    with x2:
        st.markdown("**Promoter / Governance Risk Indicator**")
        st.markdown(_status_chip("Governance Signal", str(initial_doc_signals["governance_mentions"])), unsafe_allow_html=True)

    st.markdown("**Document Processing Details (Native vs OCR Fallback)**")
    st.dataframe(pd.DataFrame(unstructured_processing), use_container_width=True)

    st.markdown("**Source-Wise Findings**")
    for title, line in source_findings.items():
        st.markdown(f"- **{title}**: {line}")

with tab3:
    st.subheader("GST vs Bank Reconciliation and Detected Flags")
    st.dataframe(recon["table"], use_container_width=True)
    r1, r2, r3 = st.columns(3)
    r1.metric("GST Turnover", f"INR {recon['gst_turnover']:,.0f}")
    r2.metric("Bank Inflows", f"INR {recon['bank_inflows']:,.0f}")
    r3.metric("Variance", f"{recon['variance_pct']:.2f}%")
    f1, f2, f3 = st.columns(3)
    with f1:
        st.markdown(_status_chip("Mismatch", recon["mismatch_flag"]), unsafe_allow_html=True)
    with f2:
        st.markdown(_status_chip("Circular Trading", recon["circular_flag"]), unsafe_allow_html=True)
    with f3:
        st.markdown(_status_chip("Revenue Inflation", recon["inflated_revenue_flag"]), unsafe_allow_html=True)

    chart_df = bank_df.copy()
    if "date" in chart_df.columns:
        chart_df["month"] = chart_df["date"].dt.to_period("M").astype(str)
        month_flow = chart_df.groupby("month", as_index=True)[["inflow", "outflow"]].sum()
        st.markdown("**Monthly Bank Flow Trend**")
        st.line_chart(month_flow)

    gst_chart = gst_df[["month", "taxable_turnover"]].set_index("month")
    st.markdown("**GST Turnover Trend**")
    st.bar_chart(gst_chart)

with tab4:
    st.subheader("Research Insights and Qualitative Findings")
    st.info(insights["ai_summary"])
    i1, i2 = st.columns(2)
    with i1:
        st.markdown(_status_chip("Regulatory Risk", insights["regulatory_risk"]), unsafe_allow_html=True)
        st.caption(insights["sector_headwind"])
    with i2:
        st.markdown(_status_chip("Promoter News Level", insights["promoter_news_level"]), unsafe_allow_html=True)
        st.caption(insights["promoter_news"])
    st.markdown("**Primary Due Diligence Notes**")
    st.caption(primary_notes if primary_notes else "Not provided")
    st.markdown("**External Intelligence Notes**")
    st.caption(external_notes if external_notes else "Not provided")

with tab5:
    st.subheader("Risk Scoring and Explainability Dashboard")
    score_df = pd.DataFrame([{"Five C": k, "Score (out of 20)": v} for k, v in score_bundle["five_cs_scores"].items()])
    st.dataframe(score_df, use_container_width=True)
    p1, p2 = st.columns(2)
    with p1:
        st.markdown("**Top Positive Drivers**")
        for item in score_bundle.get("positive_drivers", []):
            st.write(f"- {item}")
    with p2:
        st.markdown("**Top Negative Drivers**")
        for item in score_bundle.get("negative_drivers", []):
            st.write(f"- {item}")

    st.markdown("**Penalties and Boosts Applied**")
    adj_df = pd.DataFrame(score_bundle.get("adjustments", []))
    if not adj_df.empty:
        st.dataframe(adj_df, use_container_width=True)

    st.markdown("**RAG-Derived Evidence Signals**")
    rag_df = pd.DataFrame(rag_adjustments)
    if not rag_df.empty:
        st.dataframe(rag_df[["signal", "factor", "points", "reason", "evidence_source", "evidence_chunk"]], use_container_width=True)
    st.markdown("**Qualitative Score Adjustments**")
    qual_df = pd.DataFrame(qual_adjustments)
    if not qual_df.empty:
        st.dataframe(qual_df, use_container_width=True)
    unified_rows = []
    for row in qual_adjustments:
        unified_rows.append(
            {
                "source": row.get("source"),
                "origin": "primary_due_diligence" if row.get("source") == "Primary Due Diligence" else "external_intelligence",
                "cue": row.get("signal"),
                "five_c": row.get("factor"),
                "score_effect": row.get("points"),
                "reason": row.get("reason"),
            }
        )
    for row in rag_adjustments:
        unified_rows.append(
            {
                "source": row.get("source"),
                "origin": "unstructured_rag",
                "cue": row.get("signal"),
                "five_c": row.get("factor"),
                "score_effect": row.get("points"),
                "reason": row.get("reason"),
            }
        )
    for row in research_agent_findings.get("signal_rows", []):
        unified_rows.append(
            {
                "source": row.get("source"),
                "origin": row.get("origin"),
                "cue": row.get("cue"),
                "five_c": row.get("five_c"),
                "score_effect": row.get("score_effect"),
                "reason": row.get("reason"),
            }
        )
    st.markdown("**Unified Explainability Signals**")
    if unified_rows:
        st.dataframe(pd.DataFrame(unified_rows), use_container_width=True)
    st.markdown("**Decision, Limit, and Pricing Explainability**")
    st.success(recommendation["explanation"])
    st.write(f"- Decision logic: {recommendation['decision_rationale']}")
    st.write(f"- Limit logic: {recommendation['limit_rationale']}")
    st.write(f"- Pricing logic: {recommendation['pricing_rationale']}")
    st.markdown("**Indian-Context Indicators Used**")
    st.dataframe(pd.DataFrame(indian_context_cards), use_container_width=True)

with tab6:
    st.subheader("Final CAM Preview and Recommendation")
    c1, c2, c3 = st.columns(3)
    c1.metric("Scorecard Decision", recommendation["decision"])
    c2.metric("ML Decision", ml_view["ml_decision"])
    c3.metric("Final Hybrid Decision", hybrid_view["final_decision"])
    st.text_area("Credit Appraisal Memo", value=cam_text, height=520)
    d1, d2 = st.columns(2)
    with d1:
        st.download_button(
            "Download CAM (.txt)",
            cam_text,
            f"Intelli_Credit_CAM_{company_name.replace(' ', '_')}.txt",
            "text/plain",
            key="cam_txt_main",
        )
    with d2:
        if docx_bytes:
            st.download_button(
                "Download CAM (.docx)",
                docx_bytes,
                f"Intelli_Credit_CAM_{company_name.replace(' ', '_')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="cam_docx_main",
            )

with tab7:
    st.subheader("Unstructured Intelligence (Local RAG)")
    c1, c2, c3 = st.columns(3)
    c1.metric("Documents Indexed", rag_index.get("document_count", 0))
    c2.metric("Chunk Count", rag_index.get("chunk_count", 0))
    c3.metric("Retrieval Status", rag_index.get("status", "empty"))
    for theme, result in rag_theme_results.items():
        with st.container(border=True):
            st.markdown(f"**{theme}**")
            st.caption(result.get("summary", ""))
            st.write(f"Contributes to risk analysis: {'Yes' if result.get('contributes_to_risk') else 'No'}")
            for snip in result.get("snippets", []):
                st.markdown(f"- `{snip.get('source')} | {snip.get('chunk_id')} | score={snip.get('score')}`: {snip.get('snippet')}")

with tab8:
    st.subheader("Research Agent Panel")
    if st.button("Generate Research Summary", key="gen_research_summary"):
        st.session_state["research_summary_text"] = research_agent_findings.get("summary", "")
    if "research_summary_text" not in st.session_state:
        st.session_state["research_summary_text"] = research_agent_findings.get("summary", "")
    st.info(st.session_state.get("research_summary_text", "No research summary generated."))
    for title, value in research_agent_findings.get("cards", {}).items():
        st.markdown(f"**{title}**")
        st.caption(value)

with tab9:
    st.subheader("Indian Context Signals")
    for card in indian_context_cards:
        with st.container(border=True):
            st.markdown(_status_chip(card["Indicator"], card["Status"]), unsafe_allow_html=True)
            st.caption(f"Type: {card['Type']}")
            st.caption(card["Detail"])
            st.caption(f"Five C relevance: {card['Five C Relevance']}")

with tab10:
    st.subheader("Hybrid Recommendation")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown("**Scorecard Risk View**")
        st.write(f"- Decision: {recommendation['decision']}")
        st.write(f"- Risk Category: {score_bundle['risk_category']}")
        st.write(f"- Suggested Limit: INR {recommendation['recommended_limit']:,.0f}")
        st.write(f"- Pricing: {recommendation['interest_band']}")
    with c2:
        st.markdown("**ML Risk View**")
        st.write(f"- Decision: {ml_view['ml_decision']}")
        st.write(f"- Risk View: {ml_view['ml_risk']}")
        st.write(f"- Confidence: {ml_view['ml_confidence']}")
        st.write(f"- Probabilities: {ml_view['class_probabilities']}")

    st.markdown("**Final Hybrid Decision**")
    st.success(
        f"Final: {hybrid_view['final_decision']} | Limit: INR {hybrid_view['hybrid_limit']:,.0f} | "
        f"Pricing: {hybrid_view['hybrid_pricing_band']} | Confidence: {hybrid_view['decision_confidence']:.2f}"
    )
    st.caption(hybrid_view["alignment_note"])
    st.caption(hybrid_view["rationale"])

with tab11:
    st.subheader("ML Diagnostics")
    st.markdown("**Engineered Features Used**")
    st.dataframe(pd.DataFrame([ml_features]).T.rename(columns={0: "value"}), use_container_width=True)
    st.markdown("**Top ML Feature Importance**")
    st.dataframe(pd.DataFrame(ml_view["top_features"]), use_container_width=True)

with tab12:
    st.subheader("Export CAM")
    st.text_area("CAM Preview", value=cam_text, height=420)
    e1, e2, e3 = st.columns(3)
    with e1:
        st.download_button(
            "Download CAM (.txt)",
            cam_text,
            f"Intelli_Credit_CAM_{company_name.replace(' ', '_')}.txt",
            "text/plain",
            key="cam_txt_export",
        )
    with e2:
        if docx_bytes:
            st.download_button(
                "Download CAM (.docx)",
                docx_bytes,
                f"Intelli_Credit_CAM_{company_name.replace(' ', '_')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="cam_docx_export",
            )
        else:
            st.caption("DOCX dependency unavailable.")
    with e3:
        if pdf_bytes:
            st.download_button(
                "Download CAM (.pdf)",
                pdf_bytes,
                f"Intelli_Credit_CAM_{company_name.replace(' ', '_')}.pdf",
                "application/pdf",
                key="cam_pdf_export",
            )
        else:
            st.caption("PDF dependency unavailable.")

with tab13:
    st.subheader("Live Research (Optional)")
    st.markdown(f"**Mode:** {live_research.get('mode')} | **Status:** {live_research.get('status')}")
    st.caption(live_research.get("summary", ""))
    if live_research.get("headlines"):
        for h in live_research["headlines"]:
            st.markdown(f"- {h.get('title')} ({h.get('query')})")

with tab14:
    st.subheader("How Intelli-Credit Works")
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1:
        st.markdown("<div class='flow-box'><b>Input Documents / Notes</b></div>", unsafe_allow_html=True)
    with c2:
        st.markdown("<div class='flow-box'><b>Parsing and Extraction</b></div>", unsafe_allow_html=True)
    with c3:
        st.markdown("<div class='flow-box'><b>RAG + Research Enrichment</b></div>", unsafe_allow_html=True)
    with c4:
        st.markdown("<div class='flow-box'><b>Five Cs Scoring</b></div>", unsafe_allow_html=True)
    with c5:
        st.markdown("<div class='flow-box'><b>Recommendation Engine</b></div>", unsafe_allow_html=True)
    with c6:
        st.markdown("<div class='flow-box'><b>CAM Generation</b></div>", unsafe_allow_html=True)

st.caption("Hackathon prototype. Mock intelligence is used where live integrations are unavailable.")
